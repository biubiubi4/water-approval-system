from __future__ import annotations

from pathlib import Path
from typing import List, Sequence

try:
    from langchain.schema import Document
except ImportError:
    from langchain_core.documents import Document

try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
except ImportError:
    from langchain_text_splitters import RecursiveCharacterTextSplitter

try:
    from langchain_community.document_loaders import (
        CSVLoader,
        Docx2txtLoader,
        JSONLoader,
        PyPDFLoader,
        TextLoader,
        UnstructuredFileLoader,
        UnstructuredHTMLLoader,
        UnstructuredURLLoader,
        WebBaseLoader,
    )
except ImportError:
    from langchain.document_loaders import PyPDFLoader, Docx2txtLoader, TextLoader
    CSVLoader = None  # type: ignore[assignment]
    JSONLoader = None  # type: ignore[assignment]
    UnstructuredFileLoader = None  # type: ignore[assignment]
    UnstructuredHTMLLoader = None  # type: ignore[assignment]
    UnstructuredURLLoader = None  # type: ignore[assignment]
    WebBaseLoader = None  # type: ignore[assignment]


def _load_documents_from_loader(loader: object) -> List[Document]:
    try:
        return loader.load()  # type: ignore[attr-defined]
    except Exception as error:
        print(f"文档加载失败: {error}")
        return []


def _clean_extracted_text(text: str) -> str:
    if not text:
        return ""

    cleaned = text.replace("\u00a0", " ")
    cleaned = cleaned.replace("\u200b", "")
    cleaned = cleaned.replace("\ufeff", "")
    cleaned = cleaned.replace("\ufffd", "")
    cleaned = cleaned.replace("-\n", "")
    cleaned = cleaned.replace("\r\n", "\n").replace("\r", "\n")

    lines = []
    for line in cleaned.split("\n"):
        line = " ".join(line.split())
        if line:
            lines.append(line)

    cleaned = "\n".join(lines)
    cleaned = cleaned.replace(" ,", ",").replace(" .", ".").replace(" ：", "：").replace(" ：", "：")
    cleaned = cleaned.replace(" ；", "；").replace(" ！", "！").replace(" ？", "？")
    return cleaned.strip()


def _normalize_documents(documents: List[Document], file_path: Path) -> List[Document]:
    normalized: List[Document] = []
    for document in documents:
        page_content = _clean_extracted_text(document.page_content)
        if not page_content:
            continue

        metadata = dict(document.metadata or {})
        metadata.setdefault("source", file_path.name)
        metadata.setdefault("file_path", str(file_path))
        normalized.append(Document(page_content=page_content, metadata=metadata))

    return normalized


def _create_pdf_loader(file_path: Path, extract_images: bool = False) -> object:
    loader_kwargs = {}
    if extract_images:
        loader_kwargs["extract_images"] = True

    try:
        return PyPDFLoader(str(file_path), **loader_kwargs)
    except TypeError:
        return PyPDFLoader(str(file_path))


def _load_docx_with_python_docx(file_path: Path) -> List[Document]:
    try:
        from docx import Document as WordDocument
    except Exception as error:
        print(f"python-docx 不可用 {file_path}: {error}")
        return []

    try:
        document = WordDocument(str(file_path))
    except Exception as error:
        print(f"Word 备用解析失败 {file_path}: {error}")
        return []

    lines: List[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append("\t".join(cells))

    for section in document.sections:
        header = section.header
        footer = section.footer
        for paragraph in header.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)
        for paragraph in footer.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)

    text = "\n".join(lines).strip()
    if not text:
        return []

    return [Document(page_content=text, metadata={"source": file_path.name, "file_path": str(file_path)})]


def split_text(
    text: str,
    chunk_size: int = 500,
    chunk_overlap: int = 100,
) -> List[str]:
    """将文本分割成小块"""
    text = _clean_extracted_text(text)
    if not text.strip():
        return []
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", "。", "！", "？", " ", ""],
    )
    return text_splitter.split_text(text)


def load_pdf(file_path: Path) -> List[Document]:
    """加载PDF文档"""
    try:
        loader = _create_pdf_loader(file_path, extract_images=True)
        documents = _normalize_documents(loader.load(), file_path)
        if documents:
            return documents

        fallback_loader = _create_pdf_loader(file_path, extract_images=False)
        return _normalize_documents(fallback_loader.load(), file_path)
    except Exception as e:
        print(f"PDF加载失败 {file_path}: {e}")
        return []


def load_docx(file_path: Path) -> List[Document]:
    """加载Word文档"""
    try:
        loader = Docx2txtLoader(str(file_path))
        documents = loader.load()
        if documents and any(document.page_content.strip() for document in documents):
            return _normalize_documents(documents, file_path)
        return _load_docx_with_python_docx(file_path)
    except Exception as e:
        print(f"Word加载失败 {file_path}: {e}")
        return _load_docx_with_python_docx(file_path)


def load_txt(file_path: Path) -> List[Document]:
    """加载文本文件"""
    try:
        loader = TextLoader(str(file_path), encoding="utf-8")
        return _normalize_documents(loader.load(), file_path)
    except UnicodeDecodeError:
        try:
            loader = TextLoader(str(file_path), encoding="gbk")
            return _normalize_documents(loader.load(), file_path)
        except Exception as e:
            print(f"Text加载失败 {file_path}: {e}")
            return []


def load_csv(file_path: Path) -> List[Document]:
    if CSVLoader is None:
        return load_txt(file_path)
    try:
        loader = CSVLoader(str(file_path), encoding="utf-8")
        return _normalize_documents(loader.load(), file_path)
    except Exception as error:
        print(f"CSV加载失败 {file_path}: {error}")
        return load_txt(file_path)


def load_html(file_path: Path) -> List[Document]:
    if UnstructuredHTMLLoader is None:
        return load_txt(file_path)
    try:
        loader = UnstructuredHTMLLoader(str(file_path))
        return _normalize_documents(loader.load(), file_path)
    except Exception as error:
        print(f"HTML加载失败 {file_path}: {error}")
        return load_txt(file_path)


def load_json(file_path: Path) -> List[Document]:
    if JSONLoader is None:
        return load_txt(file_path)
    try:
        loader = JSONLoader(str(file_path), jq_schema=".", text_content=False)
        return _normalize_documents(loader.load(), file_path)
    except Exception as error:
        print(f"JSON加载失败 {file_path}: {error}")
        return load_txt(file_path)


def load_generic_file(file_path: Path) -> List[Document]:
    if UnstructuredFileLoader is None:
        print(f"不支持的文件类型或未安装 unstructured 依赖: {file_path.suffix}")
        return []
    try:
        loader = UnstructuredFileLoader(str(file_path))
        return _normalize_documents(loader.load(), file_path)
    except Exception as error:
        print(f"通用文件加载失败 {file_path}: {error}")
        return load_txt(file_path)


def load_url(url: str) -> List[Document]:
    loaders: Sequence[object] = []
    if UnstructuredURLLoader is not None:
        loaders = [UnstructuredURLLoader([url])]
    elif WebBaseLoader is not None:
        loaders = [WebBaseLoader(url)]

    for loader in loaders:
        documents = _load_documents_from_loader(loader)
        if documents:
            return [Document(page_content=_clean_extracted_text(doc.page_content), metadata=dict(doc.metadata or {})) for doc in documents if _clean_extracted_text(doc.page_content)]

    return []


def load_document(file_path: Path) -> List[Document]:
    """根据文件类型加载文档"""
    suffix = file_path.suffix.lower()
    
    if suffix == ".pdf":
        return load_pdf(file_path)
    elif suffix in [".docx", ".doc"]:
        return load_docx(file_path)
    elif suffix in [".txt", ".md"]:
        return load_txt(file_path)
    elif suffix == ".csv":
        return load_csv(file_path)
    elif suffix in [".html", ".htm"]:
        return load_html(file_path)
    elif suffix == ".json":
        return load_json(file_path)
    else:
        return load_generic_file(file_path)


def process_documents(
    file_paths: List[Path],
    chunk_size: int = 500,
    chunk_overlap: int = 100,
) -> List[Document]:
    """处理文档，加载并分块"""
    all_documents: List[Document] = []
    
    for file_path in file_paths:
        if not file_path.exists():
            print(f"文件不存在: {file_path}")
            continue
        
        print(f"正在处理: {file_path.name}")
        docs = load_document(file_path)
        
        if not docs:
            continue
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", "。", "！", "？", " ", ""],
        )
        
        split_docs = text_splitter.split_documents(docs)
        
        for chunk_index, doc in enumerate(split_docs, 1):
            doc.metadata["source"] = file_path.name
            doc.metadata["file_path"] = str(file_path)
            doc.metadata["chunk_index"] = chunk_index
            doc.metadata.setdefault("document_name", file_path.stem)
        
        all_documents.extend(split_docs)
    
    print(f"处理完成，共 {len(all_documents)} 个文档块")
    return all_documents
